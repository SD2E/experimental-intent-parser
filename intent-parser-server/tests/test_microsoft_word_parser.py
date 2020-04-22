from docx import Document
from microsoft_word_parser import MicrosoftWordParser
import os
import unittest


class MicrosoftWordParserTest(unittest.TestCase):


    def setUp(self):
        curr_path = os.path.dirname(os.path.realpath(__file__))
        self.file_name = 'Nick Copy of CP Experimental Request - NovelChassisYeastStates_TimeSeries'
        file = '%s.docx' % self.file_name
        file_path = os.path.join(curr_path, 'data', file)
        document = Document(file_path)
        self.doc_parser = MicrosoftWordParser(document)

    def tearDown(self):
        pass


    def test_parse_document_title(self):
        document_title = self.doc_parser.title()
        
    def test_get_current_revision(self):
        self.assertEquals(self.doc_parser.revision(), 1)
        
    def test_getting_paragraphs(self):
        list_of_paragraphs = self.doc_parser.paragraphs()
    
    def test_tables(self):
        self.doc_parser.tables() 
        
    def test_link_info(self):
        linked_words = self.doc_parser.link_info()
        for word in linked_words:
            print('%s : %s' % (word, linked_words[word]))


if __name__ == "__main__":
    unittest.main()
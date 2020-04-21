from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.constants import NAMESPACE 
import logging

class MicrosoftWordParser(object):
    '''
    Processes content from a Microsoft Word document (.docx) 
    '''
    
    PARAGRAPH_RUN_TAG_NAME = '{%s}r' % NAMESPACE.WML_MAIN
    RUN_TEXT = '{%s}t' % NAMESPACE.WML_MAIN
    _logger = logging.getLogger('word_parser')
    
    def __init__(self, docx):
        self.document = docx
        self.core_properties = docx.core_properties
   
    def revision(self):
        """
        Get this document's revision
        
        Return:
            An integer value to represent the revision number
        """
        return self.core_properties.revision
    
    
    def link_info(self): 
        """
        Get words with hyperlinks
        """
        return self.get_hyperlinks_from_document()
    
    def paragraphs(self):
        """
        Get document paragraphs
        """
        list_of_paragraphs = self.document.paragraphs
        index = 1
        result = []
        for paragraph in list_of_paragraphs:
            text = paragraph.text
            print('%d : %s' %(index, text))
            index = index + 1  
            result.append(text)    
        return [paragraph.text for paragraph in list_of_paragraphs]
    
    def title(self):
        """
        Get title of document
        
        Returns:
            The document file name as a string
        """
        return self.core_properties.title
    
    def tables(self):
        """
        Parse all table in document
        """
        for table in self.document.tables:
            self._parse_table(table)
            
                        
    def _parse_table(self, table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    print(paragraph.text)
    
    def get_hyperlinks_from_document(self):
        rIds_to_urls = self._get_hyperlink_rIds_with_url()
        rIds_to_text = {}
        for paragraph in self.document.paragraphs:
            hyperlink_elements = paragraph._p.xpath("w:hyperlink")
            if(len(hyperlink_elements) > 0):
                for element in hyperlink_elements:
                    attributes = [r_id for _, r_id in element.items()]
                    run_texts = self.get_run_text(element)
                    if len(attributes) != 1:
                        self._logger.log('Unable to map %d rIds to a hyperlink text.' % len(attributes))
                        continue
                    for text in run_texts:
                        rIds_to_text[attributes[0]] = text
                        
        hyperlink_mapping = {}
        for rId in rIds_to_text:
            if rId in rIds_to_urls:
                hyperlink_mapping[rIds_to_text[rId]] = rIds_to_urls[rId]
        return hyperlink_mapping
                        
    def get_run_text(self, paragraph_run):
        run_texts = []
        
        for run in paragraph_run.iterchildren(tag=self.PARAGRAPH_RUN_TAG_NAME):
            for run_text in run.iterchildren(tag=self.RUN_TEXT):
                run_texts.append(run_text.text)
        return run_texts
                                    
                                    
    def _get_hyperlink_rIds_with_url(self):
        hyperlink_ids = {}
        rels = self.document.part.rels
        for rel in rels:
            if rels[rel].reltype == RT.HYPERLINK:
                hyperlink_ids[rel] = rels[rel]._target 
        return hyperlink_ids
        